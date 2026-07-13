from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
LEGAL_DIR = ROOT / "docs" / "legal"
OUT_DIR = LEGAL_DIR / "finali"
LOGO = ROOT / "public" / "brand" / "vygo-logo-horizontal.png"

CYAN = "13C7DF"
NAVY = "12232C"
CHARCOAL = "46555D"
LIGHT_CYAN = "E8FBFE"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D9E1E7"
TEXT = "1C262C"

SOURCE_FILES = [
    "01-contratto-saas-b2b-vygo.md",
    "02-condizioni-generali-vygo.md",
    "03-dpa-vygo.md",
    "04-ordine-commerciale-vygo.md",
    "05-accordo-pilota-vygo.md",
    "06-checklist-legale-vygo.md",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GRAY, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for key, value in values.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    set_run_font(run, size=8.5, color=CHARCOAL)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_end)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 15, NAVY, 14, 7),
        ("Heading 2", 12.5, CHARCOAL, 10, 5),
        ("Heading 3", 11.5, CHARCOAL, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.4)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.15)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12


def configure_page(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        hp.add_run().add_picture(str(LOGO), width=Inches(1.35))
    r = hp.add_run("   " + title)
    set_run_font(r, size=8.5, color=CHARCOAL, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    left = fp.add_run("Vygo - documento riservato - bozza da validare legalmente")
    set_run_font(left, size=8.2, color=CHARCOAL)
    fp.add_run("\t")
    add_page_number(fp)


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(10)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(2.45))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=20, color=NAVY, bold=True)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(12)
    sr = s.add_run(subtitle)
    set_run_font(sr, size=10.5, color=CHARCOAL, bold=True)

    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = 1
    labels = [
        ("VERSIONE", "Bozza operativa"),
        ("USO", "Revisione legale"),
        ("DATA", date.today().strftime("%d/%m/%Y")),
        ("STATO", "Non definitiva"),
    ]
    for cell, (label, value) in zip(meta.rows[0].cells, labels):
        set_cell_shading(cell, LIGHT_CYAN)
        set_cell_border(cell, "BFEFF6", "8")
        set_cell_margins(cell, 120, 130, 120, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(label + "\n")
        set_run_font(lr, size=7.4, color=CHARCOAL, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=8.6, color=NAVY, bold=True)

    doc.add_paragraph()


def add_callout(doc: Document, text: str, fill=LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 1
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "C9D3DA", "8")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=NAVY, bold=True)
    doc.add_paragraph()


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        # Make placeholders visually obvious.
        segments = re.split(r"(\[[^\]]+\])", clean)
        for segment in segments:
            if not segment:
                continue
            run = paragraph.add_run(segment)
            is_placeholder = segment.startswith("[") and segment.endswith("]")
            set_run_font(
                run,
                size=10.5,
                color=CHARCOAL if is_placeholder else TEXT,
                bold=bold,
                italic=is_placeholder,
            )


def add_signature_table(doc: Document) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = 1
    headers = ["Cliente", "Fornitore"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_CYAN)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run_font(r, size=10, color=NAVY, bold=True)
    for row_idx, label in enumerate(["Nome e qualifica", "Data", "Firma"], start=1):
        for col in range(2):
            cell = table.cell(row_idx, col)
            set_cell_border(cell)
            set_cell_margins(cell, 180, 160, 180, 160)
            p = cell.paragraphs[0]
            r = p.add_run(label + ": ")
            set_run_font(r, size=8.5, color=CHARCOAL, bold=True)
            line = p.add_run("_" * 34)
            set_run_font(line, size=9, color=CHARCOAL)


def parse_markdown(md_text: str):
    blocks = []
    paragraph = []

    def flush():
        nonlocal paragraph
        if paragraph:
            blocks.append(("p", " ".join(paragraph)))
            paragraph = []

    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            blocks.append((f"h{level}", line[level:].strip()))
        elif line.startswith("- [ ]"):
            flush()
            blocks.append(("check", line[5:].strip()))
        elif line.startswith("- "):
            flush()
            blocks.append(("bullet", line[2:].strip()))
        else:
            paragraph.append(line)
    flush()
    return blocks


def render_blocks(doc: Document, blocks) -> str:
    title = ""
    seen_signature_heading = False
    for kind, text in blocks:
        if kind == "h1":
            title = text
            continue
        if kind == "h2":
            heading_key = re.sub(r"^\d+\.\s*", "", text.lower().strip())
            seen_signature_heading = heading_key in {"firma", "firme"}
            doc.add_paragraph(text, style="Heading 1")
            continue
        if kind == "h3":
            doc.add_paragraph(text, style="Heading 2")
            continue
        if kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            continue
        if kind == "check":
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, "Da compilare: " + text)
            continue
        if kind == "p":
            if text.upper().startswith("BOZZA"):
                add_callout(doc, text, fill=LIGHT_CYAN)
                continue
            label_only = text in {"Cliente:", "Fornitore:", "Data:", "Firma:"}
            if seen_signature_heading and label_only:
                continue
            p = doc.add_paragraph()
            add_inline_runs(p, text)
    return title


def build_doc(source: Path) -> Path:
    md = source.read_text(encoding="utf-8")
    blocks = parse_markdown(md)
    title = next((text for kind, text in blocks if kind == "h1"), source.stem)

    doc = Document()
    configure_styles(doc)
    configure_page(doc, title)
    add_title_block(doc, title, "Pacchetto contrattuale Vygo")
    add_callout(
        doc,
        "Documento preparatorio: usare per revisione con avvocato, consulente privacy o cliente pilota. Non usare come parere legale definitivo.",
    )
    render_blocks(doc, blocks)
    if any(kind == "h2" and re.sub(r"^\d+\.\s*", "", text.lower().strip()) in {"firma", "firme"} for kind, text in blocks):
        add_signature_table(doc)

    out = OUT_DIR / (source.stem + ".docx")
    doc.save(out)
    return out


def build_index(generated: list[Path]) -> Path:
    doc = Document()
    configure_styles(doc)
    configure_page(doc, "Indice pacchetto legale Vygo")
    add_title_block(doc, "Pacchetto legale Vygo", "Documenti per pilota, clienti paganti e revisione professionale")
    add_callout(
        doc,
        "Questo pacchetto raccoglie bozze operative. Prima dell'uso commerciale vanno validate da professionisti abilitati.",
        fill=LIGHT_CYAN,
    )
    doc.add_paragraph("Documenti inclusi", style="Heading 1")
    for item in generated:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item.name)
        set_run_font(r, size=10.5, color=TEXT, bold=True)
    doc.add_paragraph("Uso consigliato", style="Heading 1")
    for line in [
        "Accordo pilota per test gratuiti o agevolati.",
        "Ordine commerciale + contratto SaaS + condizioni generali per clienti paganti.",
        "DPA privacy da allegare quando vengono trattati dati personali dei dipendenti del cliente.",
        "Checklist legale da consegnare ad avvocato e consulente privacy.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_runs(p, line)
    out = OUT_DIR / "00-indice-pacchetto-legale-vygo.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [build_doc(LEGAL_DIR / filename) for filename in SOURCE_FILES]
    build_index(generated)
    print("Generated:")
    for path in sorted(OUT_DIR.glob("*.docx")):
        print(path)


if __name__ == "__main__":
    main()
